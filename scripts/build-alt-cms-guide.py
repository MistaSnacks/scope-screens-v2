# scripts/build-alt-cms-guide.py
# Generates the Scope Screenings ALT (multipage) CMS editor guide as a .docx,
# brand-styled (Aachen Bold headings, Libre Franklin body) with an annotated
# screenshot per section.  Screenshots come from .shots/guide/alt/*.png
# (see .shots/capture-alt-cms-guide.mjs).
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SHOTS = os.path.join(_ROOT, ".shots/guide/alt")

RUST = RGBColor(0xB1, 0x3A, 0x2A)
INK = RGBColor(0x14, 0x12, 0x10)
GREY = RGBColor(0x55, 0x50, 0x4A)

_BM = {"id": 0}


def slugify(text):
    """Bookmark-safe anchor: letters/digits/underscores, starts with a letter,
    max 40 chars. Stable for a given heading string so the ToC links line up."""
    s = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    if not s:
        s = "h"
    if not s[0].isalpha():
        s = "h_" + s
    return s[:40]


def add_bookmark(paragraph, name):
    """Wrap a heading paragraph in a Word bookmark so links can jump to it."""
    p = paragraph._p
    _BM["id"] += 1
    bid = str(_BM["id"])
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    pPr = p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(start)
    else:
        p.insert(0, start)
    p.append(end)


def add_internal_link(paragraph, text, anchor, size=10.5, color=RUST, bold=False):
    """Append a clickable in-document hyperlink (jumps to a bookmark anchor)."""
    h = OxmlElement("w:hyperlink")
    h.set(qn("w:anchor"), anchor)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Libre Franklin")
    rFonts.set(qn("w:hAnsi"), "Libre Franklin")
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), str(color))
    rPr.append(col)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    h.append(r)
    paragraph._p.append(h)


def style_run(r, font="Libre Franklin", size=10.5, bold=False, color=INK):
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color


def heading(doc, text, size=14, level=1, nav=True):
    p = doc.add_paragraph()
    if nav:
        # Built-in Heading style gives the paragraph an outline level, which is
        # what populates Word's left-hand Navigation pane. Run-level formatting
        # below overrides the style's font/colour so the brand look is kept.
        try:
            p.style = doc.styles["Heading %d" % level]
        except KeyError:
            pass
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run(text), font="Aachen Bold", size=size, color=RUST)
    if nav:
        add_bookmark(p, slugify(text))
    return p


def para(doc, text, color=INK, size=10.5, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    style_run(p.add_run(text), size=size, color=color)
    return p


def bullet(doc, text, label=None):
    p = doc.add_paragraph(style="List Bullet")
    if label:
        style_run(p.add_run(label + ": "), bold=True)
    style_run(p.add_run(text))


def legend_line(doc, n, name, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    style_run(p.add_run(f"{n}  "), bold=True, color=RUST, size=11)
    style_run(p.add_run(name), bold=True)
    style_run(p.add_run(" - " + desc))


def optional_image(doc, filename, caption=None, missing_note=None):
    """Embed .shots/guide/alt/<filename> if present; otherwise leave a labelled
    placeholder so the slot is obvious until the screenshot is dropped in."""
    img = os.path.join(SHOTS, filename)
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].paragraph_format.space_after = Pt(3)
        if caption:
            c = doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.space_after = Pt(6)
            style_run(c.add_run(caption), size=9, color=GREY)
    elif missing_note:
        n = doc.add_paragraph()
        n.paragraph_format.space_after = Pt(6)
        style_run(n.add_run(missing_note), size=9, color=GREY)


# One entry per section documented — each gets its own page in the guide.
# (key, title, collections_note, [(n, name, desc), ...], also_note)
SECTIONS = [
    # 1 Hero
    ("Hero", "Hero", "Form: Hero", [
        ("1", "Eyebrow", "small label above the curtain (e.g. “Lexscope Presents”)."),
        ("2", "Video reel area", "the background reel. Poster image and Video are Media-Manager fields — open the Hero form and upload/select there."),
    ], "Also in the Hero form: Poster image (the still shown before the video plays) and Video (the background reel). "
       "The big “Scope / Screenings” wordmark and tagline are fixed in the design and are not editable here."),

    # 2 About -- header
    ("AboutHeader", "About — Page Header", "Form: AboutPage", [
        ("1", "Eyebrow", "small label above the title (e.g. “About the Festival”)."),
        ("2", "Title", "the main page heading. Use a line break in the field for two lines."),
        ("3", "Lede", "the introductory paragraph under the title."),
    ], "To hide this page, use the SiteSettings form (see “Show or hide a page”)."),

    # 3 About -- Timeline
    ("Timeline", "About — Timeline", "List: Timeline", [
        ("1", "Year", "the large year label on the left of the row."),
        ("2", "Title", "the milestone heading."),
        ("3", "Blurb", "the description paragraph."),
    ], "Order: lower numbers show first. Add a new row for each additional milestone. "
       "Note: the Timeline section is currently turned off on the live About page, so edits here "
       "won’t appear on the site until it’s switched back on."),

    # 5 Founder band
    ("FounderBand", "Founder Band", "Form: BuiltForAccess", [
        ("1", "Title", "the section heading (e.g. “Built For Access”). Line break OK."),
        ("2", "Quote", "the large blockquote."),
        ("3", "Founder name", "shown on the photo frame and beside the quote."),
        ("4", "Founder title · credential", "the role and credential line under the name."),
        ("5", "Photo", "upload or select the founder portrait in Wix Media Manager."),
    ], "This band appears on the homepage and on /about — edit it once and both pages update."),

    # 6 Schedule
    ("Schedule", "Schedule Page", "Form: SchedulePage", [
        ("1", "Eyebrow", "small label above the title (e.g. “The Season”)."),
        ("2", "Title", "the main page heading."),
        ("3", "Lede", "the introductory paragraph."),
        ("4", "Closing band", "the heading, body copy, button text and href at the bottom of the page."),
    ], "To hide this page, use the SiteSettings form (see “Show or hide a page”). The schedule grid and event dates come from "
       "Wix Events — the CMS only controls the surrounding words."),

    # 7 Tickets
    ("Tickets", "Tickets Page", "Form: TicketsPage", [
        ("1", "Eyebrow", "small label above the title (e.g. “Chapter One”)."),
        ("2", "Title", "the main page heading."),
        ("3", "Lede", "the introductory paragraph."),
        ("4", "Why-a-season-pass title", "the heading of the season-pass explainer block."),
        ("5", "Why body", "the body copy in the season-pass block."),
        ("6", "Closing band", "the heading, body copy, button text and href at the bottom of the page."),
    ], "To hide this page, use the SiteSettings form (see “Show or hide a page”). Ticket checkout is handled by Wix Events — "
       "the CMS only controls the words around it."),

    # 8 Submit -- header
    ("SubmitHeader", "Submit Page — Header", "Form: SubmitPage", [
        ("1", "Eyebrow", "small label above the title (e.g. “Open call · via FilmFreeway”)."),
        ("2", "Title", "the main page heading."),
        ("3", "Lede", "the introductory paragraph."),
        ("4", "Closing band", "the heading, body copy, button text and href at the bottom of the page."),
    ], "To hide this page, use the SiteSettings form (see “Show or hide a page”)."),

    # 9 Submit -- Criteria
    ("SubmitCriteria", "Submit — What We’re After", "List: SubmitCriteria", [
        ("1", "Number", "the large display number on the card (e.g. 01)."),
        ("2", "Title", "the criterion heading."),
        ("3", "Blurb", "the description paragraph."),
    ], "Order: lower numbers show first. One row per criterion card."),

    # 10 Submit -- Steps
    ("SubmitSteps", "Submit — Three Steps In", "List: SubmitSteps", [
        ("1", "Number", "the large display number on the card (e.g. 01)."),
        ("2", "Title", "the step heading."),
        ("3", "Blurb", "the description paragraph."),
    ], "Order: lower numbers show first. One row per step card."),

    # 11 Submit -- Deadlines
    ("SubmitDeadlines", "Submit — Mark Your Calendar", "List: SubmitDeadlines", [
        ("1", "Round name", "the deadline label (e.g. EarlyScope, Regular)."),
        ("2", "Closes", "the closing date as display text (e.g. Aug 31)."),
        ("3", "Fee", "the submission fee as display text (e.g. $15)."),
    ], "Order: lower numbers show first. Update this list to match your actual FilmFreeway dates."),

    # 12 Support -- header
    ("SupportHeader", "Support Page — Header", "Form: SupportPage", [
        ("1", "Eyebrow", "small label above the title (e.g. “Funders & philanthropy”)."),
        ("2", "Title", "the main page heading."),
        ("3", "Lede", "the introductory paragraph."),
        ("4", "Give card (label / body / donate URL)", "the small card in the header area — label, body copy, and where the donate button links."),
        ("5", "Closing band", "the heading, body copy, button text and href at the bottom of the page."),
    ], "To hide this page, use the SiteSettings form (see “Show or hide a page”)."),

    # 13 Support -- Giving tiers
    ("GivingTiers", "Support — Pick Your Level", "List: GivingTiers", [
        ("1", "Name", "the tier name (e.g. Friend, Supporter)."),
        ("2", "Amount", "the dollar amount displayed on the card (e.g. $50)."),
        ("3", "Cadence", "the billing cadence (e.g. /year, /month)."),
        ("4", "Perks", "the perk lines — one perk per line inside the field."),
        ("5", "Featured highlight", "check to give this tier a rust-coloured border and ring."),
    ], "Order: lower numbers show first."),

    # 14 Footer
    ("Footer", "Footer", "Form: Footer", [
        ("1", "Sign-off", "the big closing heading in the footer (e.g. “See You At The Movies”)."),
        ("2", "Newsletter heading", "the text above the email sign-up field."),
        ("3", "Tagline", "the short paragraph under the Scope Screenings name column."),
        ("4", "Copyright", "the small legal line at the very bottom."),
        ("5", "Contact email", "the contact address shown in the footer links."),
    ], "The Footer form is shared across every page — edit it once and it updates everywhere."),

    # 15 Nav
    ("Nav", "Navigation", "List: Nav", [
        ("1", "Watch", "the first menu item label and href."),
        ("2", "Tickets", "links to the /tickets page."),
        ("3", "Schedule", "links to the /schedule page."),
        ("4", "Submit", "links to the /submit page."),
        ("5", "About", "links to the /about page."),
        ("6", "Support", "links to the /support page."),
    ], "Each row has: Label (the link text), Href (the path), Order (lower shows first), and Hidden "
       "(check to remove from the menu without deleting the row). If you hide a page from the SiteSettings "
       "form, its menu link stays in the menu but automatically points to that section on the homepage."),

    # 16 Settings — show/hide switches
    ("Settings", "Settings — On/Off Switches", "Form: SiteSettings", [
        ("1", "Hide About page", "turns the /about page off; its menu link then points to the About section on the homepage."),
        ("2", "Hide Schedule page", "turns the /schedule page off; menu link points to the homepage Schedule section."),
        ("3", "Hide Tickets page", "turns the /tickets page off; menu link points to the homepage Tickets section."),
        ("4", "Hide Submit page", "turns the /submit page off; menu link points to the homepage Submit section."),
        ("5", "Hide Support page", "turns the /support page off; menu link points to the homepage Support section."),
        ("6", "Hide Press kit", "hides the Press & Media (press kit) card in the homepage support section."),
        ("7", "Hide Archives section", "hides “The Archives” (Chapter Four) section on the homepage."),
        ("8", "Hide newsletter band", "hides the footer sign-off + email signup band."),
    ], "All of the site’s on/off switches live on this one form, and each checkbox has help text beneath it "
       "describing exactly what it hides. Check a box and Publish to hide; uncheck and Publish to show. Turning "
       "a page off never leaves a dead link — the menu item routes to the matching homepage section instead."),
]


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.7)

    # Cover header
    logo = os.path.join(_ROOT, "public/popcorn-logo.png")
    if os.path.exists(logo):
        doc.add_picture(logo, height=Inches(0.7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(h.add_run("Editing Your Website Content"), font="Aachen Bold", size=24, color=RUST)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(sub.add_run("A visual guide to the Scope Screenings CMS (multipage site)"), size=12, color=GREY)

    # --- Clickable table of contents (and it drives Word's left-hand Navigation pane) ---
    doc.add_page_break()
    toc_h = doc.add_paragraph()
    toc_h.paragraph_format.space_after = Pt(4)
    style_run(toc_h.add_run("Contents"), font="Aachen Bold", size=18, color=RUST)
    para(doc, "Open Word’s Navigation pane (View → Navigation pane on Windows, or View → Sidebar → Navigation "
              "on Mac) for a clickable outline of this guide down the left-hand side. Or click any line below "
              "to jump straight to that part.", color=GREY, size=9.5)

    def toc_group(label):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        style_run(p.add_run(label), font="Aachen Bold", size=11, color=INK)

    def toc_item(label):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.28)
        add_internal_link(p, label, slugify(label))

    toc_group("Getting started")
    for _lbl in ["The one rule", "Opening the CMS", "How it’s organized", "Your collections at a glance"]:
        toc_item(_lbl)
    toc_group("Editing each section")
    for _key, _title, *_rest in SECTIONS:
        toc_item(_title)
    toc_group("Managing the site")
    for _lbl in ["Show or hide a page", "Photos & video", "Leave these alone (they update themselves)", "Good to know"]:
        toc_item(_lbl)
    doc.add_page_break()

    # The one rule
    heading(doc, "The one rule")
    para(doc, "Open the CMS in your Wix dashboard → pick the collection for the part of the page you want to "
              "change → edit the fields → click Publish. Your live site updates within about 10 seconds.")

    # Opening the CMS
    heading(doc, "Opening the CMS")
    bullet(doc, "Sign in at wix.com and open this site’s Dashboard.", label="1")
    bullet(doc, "In the left-hand menu choose CMS. (On some plans it sits under “Dev Mode” or "
                "“Developer Tools”, and it was previously called “Content Manager”.)", label="2")
    bullet(doc, "You’ll land on the list of collections — one folder per part of the site. Click any "
                "collection to open it.", label="3")
    bullet(doc, "Edit the fields, then click Publish (top-right) to push it live.", label="4")
    optional_image(doc, "cms-nav.png",
                   caption="Where to find CMS in the Wix dashboard menu.",
                   missing_note="[ Screenshot to add → save as .shots/guide/alt/cms-nav.png: the Wix dashboard "
                                "left menu with “CMS” highlighted. ]")
    optional_image(doc, "cms-collections.png",
                   caption="The collection list — one folder per section of the site.",
                   missing_note="[ Screenshot to add → save as .shots/guide/alt/cms-collections.png: the CMS "
                                "collections list. ]")

    # How it's organized
    heading(doc, "How it’s organized")
    para(doc, "This is the multipage version of the site. Each page (/about, /schedule, /tickets, /submit, /support) "
              "has its own Page Form that controls its header text and closing call-to-action. All the show/hide "
              "switches live together on one SiteSettings form. Three forms are shared across every page: Hero (the "
              "homepage curtain), BuiltForAccess (the Founder band), and Footer. Repeating items — history, giving "
              "tiers, nav links, deadlines, and so on — live in their own lists, one row per item. On each "
              "screenshot below the numbered pins show exactly what each field controls.")

    # Collections at a glance
    heading(doc, "Your collections at a glance")
    para(doc, "Each item below is one collection (a “folder”) in the CMS. Page forms are a single editable "
              "card; lists hold multiple rows.")
    bullet(doc, "the homepage opening curtain — eyebrow, poster image, background video.", label="Hero")
    bullet(doc, "the Founder band — quote, name, title, photo. Appears on homepage and /about.", label="BuiltForAccess")
    bullet(doc, "the site footer on every page — sign-off, newsletter heading, tagline, copyright, contact.", label="Footer")
    bullet(doc, "the /about page header text.", label="AboutPage")
    bullet(doc, "the /about history items (list — one row per milestone). Currently turned off on the live site.", label="Timeline")
    bullet(doc, "the /schedule page header, closing band, and Hidden toggle.", label="SchedulePage")
    bullet(doc, "the /tickets page header, season-pass block, closing band, and Hidden toggle.", label="TicketsPage")
    bullet(doc, "the /submit page header, closing band, and Hidden toggle.", label="SubmitPage")
    bullet(doc, "“What We’re After” cards on /submit (list — one row per criterion).", label="SubmitCriteria")
    bullet(doc, "“Three Steps In” cards on /submit (list — one row per step).", label="SubmitSteps")
    bullet(doc, "the deadlines table on /submit (list — one row per deadline).", label="SubmitDeadlines")
    bullet(doc, "the /support page header, Give card, closing band, and Hidden toggle.", label="SupportPage")
    bullet(doc, "the /support giving levels (list — one row per tier).", label="GivingTiers")
    bullet(doc, "the top navigation menu — label, href, order, hidden (list).", label="Nav")
    bullet(doc, "every on/off switch in one place — each page, plus the press kit, the Archives section, and the newsletter band (form — checkboxes, each with help text).", label="SiteSettings")
    bullet(doc, "the scrolling marquee phrases (list — one row per phrase).", label="Marquee")

    # Per-section pages
    for key, title, collections, items, also in SECTIONS:
        doc.add_page_break()
        heading(doc, title)
        c = doc.add_paragraph()
        c.paragraph_format.space_after = Pt(6)
        style_run(c.add_run(collections), size=9.5, bold=True, color=GREY)
        img = os.path.join(SHOTS, f"{key}.png")
        if os.path.exists(img):
            doc.add_picture(img, width=Inches(7.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            style_run(p.add_run(f"[ Screenshot: save as .shots/guide/alt/{key}.png ]"), size=9, color=GREY)
        for n, name, desc in items:
            legend_line(doc, n, name, desc)
        if also:
            a = doc.add_paragraph()
            a.paragraph_format.space_before = Pt(4)
            style_run(a.add_run(also), size=10, color=GREY)

    # Show or hide a page
    doc.add_page_break()
    heading(doc, "Show or hide a page")
    para(doc, "All the show/hide switches live on one form: SiteSettings. Open it, check the box for the page "
              "you want to turn off (About, Schedule, Tickets, Submit, or Support), and click Publish.")
    para(doc, "When a page is hidden, its item in the top menu doesn’t disappear or break — it automatically "
              "points to the matching section on the homepage instead. Hide About, for example, and the About "
              "menu link sends visitors to the About section on the home page rather than the /about page.")
    bullet(doc, "Open SiteSettings → check the page’s box (e.g. “Schedule hidden”) → Publish. The page goes "
                "offline and its menu link now scrolls to that section on the homepage.", label="To hide")
    bullet(doc, "Open SiteSettings → uncheck it → Publish. The full page and its normal menu link come back.",
           label="To restore")
    para(doc, "Press kit: SiteSettings also has a “Hide Press kit” switch that hides the Press & Media card "
              "in the homepage support section — flip it the same way.")
    para(doc, "Same form, more switches: “Hide Archives section” turns off The Archives (Chapter Four) block on "
              "the homepage, and “Hide newsletter band” turns off the footer sign-off + email signup. Every "
              "switch on this form shows help text in the CMS describing exactly what it hides.")
    para(doc, "The Nav list is separate: check Hidden on a Nav row to remove that item from the menu entirely, "
              "rather than rerouting it. Use a SiteSettings page switch to reroute a link to the homepage; use "
              "the Nav row’s Hidden toggle to drop the link altogether.")

    # Photos & video
    heading(doc, "Photos & video")
    bullet(doc, "Hero Poster image and Hero Video: open the Hero form and upload or select in Wix Media Manager. "
                "Leave the field empty and the built-in default shows — nothing breaks.")
    bullet(doc, "BuiltForAccess Founder photo: same — upload in the form, or leave empty for the default.")
    bullet(doc, "Hero Video is the background reel on the homepage. Upload an MP4 or select an existing video "
                "in the Media Manager.")

    # Leave these alone
    heading(doc, "Leave these alone (they update themselves)")
    bullet(doc, "Tickets checkout and the Schedule dates come straight from Wix Events — manage those where "
                "you already sell tickets. The CMS only controls the surrounding words on those pages.")
    bullet(doc, "SubmitDeadlines controls the display copy on /submit — update it to match your actual "
                "FilmFreeway dates, but the live purchase flow is outside the CMS.")

    # Good to know
    heading(doc, "Good to know")
    bullet(doc, "Every list row has an Order number — lower numbers show first.")
    bullet(doc, "Two lines in one field (e.g. a title): press Enter inside the field to add the line break.")
    bullet(doc, "Changes appear on the live site within about 10 seconds. If you don’t see it, refresh the page.")

    out = os.path.join(_ROOT, "docs/Scope-Screenings-ALT-CMS-Guide.docx")
    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
